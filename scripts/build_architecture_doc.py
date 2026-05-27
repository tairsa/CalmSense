"""Generate ARCHITECTURE.docx describing the CalmSense system."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = REPO_ROOT / "ARCHITECTURE.docx"


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x14, 0x3D, 0x59)


def add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def build() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("CalmSense — System Architecture")
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = RGBColor(0x14, 0x3D, 0x59)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("Galaxy Watch 5 → Android phone → FastAPI backend")
    s_run.italic = True
    s_run.font.size = Pt(12)

    doc.add_paragraph()

    add_heading(doc, "1. Overview", level=1)
    add_para(
        doc,
        "CalmSense is a panic-attack detection prototype. A Galaxy Watch 5 streams "
        "heart-rate and motion data to a paired Android phone, which uploads the "
        "samples to a FastAPI backend and runs an on-device sigmoid classifier "
        "for low-latency panic detection. The backend stores samples, trains "
        "per-user logistic-regression weights, and serves them back to the phone.",
    )

    add_heading(doc, "Components at a glance", level=2)
    add_table(
        doc,
        ["Tier", "Process", "Responsibility"],
        [
            ["Wearable", "HrMonitoringService (foreground)",
             "Read raw HR + accelerometer, compute motion RMS, send sample over Wear MessageClient."],
            ["Phone", "WatchListenerService",
             "Receive /calmsense/sample messages, write into WatchVitalsRepository."],
            ["Phone", "MonitorService (foreground)",
             "Poll vitals, run sigmoid model on-device, POST to backend, fire panic notifications."],
            ["Phone", "MainActivity (Compose UI)",
             "Show live HR, server status, model status, breathing coach overlay."],
            ["Backend", "FastAPI + uvicorn",
             "Receive sensor data, store in Supabase (JSON fallback locally), return per-user weights."],
        ],
    )

    add_heading(doc, "2. Data flow", level=1)
    add_para(
        doc,
        "Each sample produced by the watch traverses four boundaries before it "
        "influences the user-visible state. The diagram below traces a single "
        "heart-rate event end to end.",
    )
    add_code(
        doc,
        "Galaxy Watch 5\n"
        "  Sensor.TYPE_HEART_RATE  ─┐\n"
        "  Sensor.TYPE_LINEAR_ACCEL ┘──► HrMonitoringService\n"
        "                                  │  CSV \"<bpm>,<rms>\"\n"
        "                                  ▼\n"
        "                          Wear MessageClient\n"
        "                                  │  path: /calmsense/sample\n"
        "                                  ▼\n"
        "Phone (Android)\n"
        "  WatchListenerService ─► WatchVitalsRepository\n"
        "                                  │  Vitals(hr, motion)\n"
        "                                  ▼\n"
        "                          MonitorService (poll loop)\n"
        "                                  ├─► PanicModel.predict()  (on-device sigmoid)\n"
        "                                  ├─► panic notification (debounced)\n"
        "                                  └─► BackendClient.postSensorData()\n"
        "                                              │  HTTP POST\n"
        "                                              ▼\n"
        "Backend (FastAPI)\n"
        "  POST /api/v1/sensor-data  ─► storage.py ─► Supabase\n"
        "  GET  /api/v1/sensor-data  ◄─ logistic-regression weights\n",
    )

    add_heading(doc, "3. Wearable module", level=1)
    add_heading(doc, "Permissions", level=2)
    add_bullets(
        doc,
        [
            "BODY_SENSORS / BODY_SENSORS_BACKGROUND — required to read the raw HR sensor.",
            "android.permission.health.READ_HEART_RATE — Galaxy Watch firmware gates the HR sensor at the sensor-service level on this permission.",
            "ACTIVITY_RECOGNITION — required as the companion permission for a FOREGROUND_SERVICE_TYPE_HEALTH service on Android 14+.",
            "FOREGROUND_SERVICE / FOREGROUND_SERVICE_HEALTH — long-running monitoring.",
            "POST_NOTIFICATIONS — for the ongoing CalmSense notification.",
        ],
    )

    add_heading(doc, "HrMonitoringService", level=2)
    add_para(
        doc,
        "A foreground service that uses raw android.hardware.SensorManager (not "
        "Wear Health Services) to collect data:",
    )
    add_bullets(
        doc,
        [
            "TYPE_HEART_RATE @ SENSOR_DELAY_NORMAL — ~1 Hz HR samples with accuracy 3 when worn; 0 with accuracy −1 when off-wrist.",
            "TYPE_LINEAR_ACCELERATION @ SENSOR_DELAY_UI — ~16 Hz; maintained as a rolling RMS over a 5-second window.",
            "Throttled to one outbound message every ~2 s; format is the ASCII CSV \"<bpm>,<rms>\".",
            "MessageClient send runs on a single-thread Executor — sensor callbacks fire on the main thread and Tasks.await() throws there.",
        ],
    )

    add_heading(doc, "Why not Wear Health Services?", level=2)
    add_para(
        doc,
        "PassiveMonitoringClient and MeasureClient were both tried first. On "
        "Galaxy Watch 5 (One UI Watch, current firmware), WHS_PermissionPolicy "
        "rejects every dispatch with SecurityException at bdk.m(PG:116) despite "
        "the health.READ_HEART_RATE permission being granted (granted=true, "
        "USER_SET). Health Connect's Jetpack client (HealthConnectClient) is "
        "not supported on Wear OS at all and throws UnsupportedOperationException. "
        "The raw sensor path is the only route that delivers HR data on this "
        "watch.",
    )

    add_heading(doc, "4. Phone module", level=1)

    add_heading(doc, "WatchListenerService", level=2)
    add_para(
        doc,
        "Receives Wear MessageClient events. Parses the CSV payload "
        "(\"<bpm>,<motion_rms>\") and updates WatchVitalsRepository. Also "
        "accepts the legacy \"<bpm>\" payload on /calmsense/hr for back-compat.",
    )

    add_heading(doc, "WatchVitalsRepository", level=2)
    add_para(
        doc,
        "An in-memory hot store of the latest watch reading. Treats samples "
        "older than 2 minutes as stale (returns nulls). Derives Vitals.isMoving "
        "from motion_intensity using a 0.5 m/s² threshold; raw motion_intensity "
        "is also exposed for the backend payload.",
    )

    add_heading(doc, "MonitorService", level=2)
    add_para(
        doc,
        "A foreground service that polls vitals every 30 s (5 s when HR is "
        "elevated), runs the panic check, and POSTs to the backend. Uses "
        "WatchVitalsRepository when a fresh watch reading exists; falls back "
        "to HealthConnectVitalsRepository otherwise.",
    )

    add_heading(doc, "MainActivity & UI", level=2)
    add_para(
        doc,
        "Compose UI that shows HR, HRV, motion status, server-status chip, "
        "model-status chip, and a breathing coach overlay backed by a TTS "
        "BreathingCoach. The default data source is now WATCH (live). The "
        "UI also exposes the legacy SIMULATED / HEALTH_CONNECT switches for "
        "demos.",
    )

    add_heading(doc, "On-device classifier (PanicModel)", level=2)
    add_para(
        doc,
        "Logistic-regression weights fetched from the backend are evaluated "
        "locally with a sigmoid over (HR, HRV, motion_intensity). The "
        "probability drives panic notifications (debounced via wasInPanic) "
        "and a p(panic) chip in the UI. Until the backend returns trained "
        "weights, the model is flagged \"untrained\" and the code falls back "
        "to the legacy rule (HR > 120 ∧ HRV < 20 ∧ ¬moving).",
    )

    add_heading(doc, "5. Backend (calmsense-backend)", level=1)
    add_para(
        doc,
        "FastAPI + uvicorn on port 8000. Two endpoints power the loop:",
    )
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ["POST /api/v1/sensor-data",
             "Receives SensorPayload (user_id, hr, hrv, motion, panic flag, timestamp)."],
            ["GET  /api/v1/sensor-data?user_id=…",
             "Returns logistic-regression weights for that user, or zero-vector defaults."],
            ["GET  /health",
             "Liveness probe — used by the phone's server-status chip."],
        ],
    )
    add_para(
        doc,
        "Storage is isolated in storage.py — primary backend is Supabase "
        "(Postgres) and a flat JSON file is used as a development fallback. "
        "User isolation is by user_id only; JWT auth is a planned next step.",
    )

    add_heading(doc, "6. Network topology", level=1)
    add_bullets(
        doc,
        [
            "Watch ↔ phone: Google Play Services Wearable Data Layer (MessageClient).",
            "Phone ↔ backend: plain HTTP over LAN. BACKEND_URL constant in MainActivity.kt currently points at http://192.168.1.72:8000. AndroidManifest sets usesCleartextTraffic=\"true\" for the dev configuration.",
            "Watch never speaks to the backend directly — all backend traffic is mediated by the phone's MonitorService.",
        ],
    )

    add_heading(doc, "7. Repository layout", level=1)
    add_code(
        doc,
        "CalmSense/\n"
        "├── android-app/\n"
        "│   ├── app/           — Phone module (com.example.app)\n"
        "│   │   └── src/main/java/com/example/app/\n"
        "│   │       ├── MainActivity.kt\n"
        "│   │       ├── MonitorService.kt\n"
        "│   │       └── data/\n"
        "│   │           ├── BackendApi.kt        ← weights fetch\n"
        "│   │           ├── BackendClient.kt     ← ping + POST\n"
        "│   │           ├── BreathingCoach.kt    ← TTS\n"
        "│   │           ├── PanicModel.kt        ← on-device sigmoid\n"
        "│   │           ├── HealthConnectVitalsRepository.kt\n"
        "│   │           ├── WatchListenerService.kt\n"
        "│   │           ├── WatchVitalsRepository.kt\n"
        "│   │           └── Vitals.kt\n"
        "│   └── wear/          — Watch module (com.example.app.wear)\n"
        "│       └── src/main/java/com/example/app/wear/\n"
        "│           ├── WearMainActivity.kt\n"
        "│           └── HrMonitoringService.kt\n"
        "└── calmsense-backend/\n"
        "    ├── main.py        — FastAPI routes\n"
        "    ├── models.py      — Pydantic schemas\n"
        "    └── storage.py     — Supabase + JSON fallback\n",
    )

    add_heading(doc, "8. Known limits & next steps", level=1)
    add_bullets(
        doc,
        [
            "ADB grant of health.READ_HEART_RATE flips the flag but does not enable sensor delivery — the user must tap Allow on the runtime dialog at least once per install.",
            "Wireless ADB on Galaxy Watch 5 drops the connection silently; expect to re-pair occasionally.",
            "BackendApi and BackendClient still coexist after the merge that brought in the Supabase / sigmoid work. Consolidating to a single client is a cleanup task.",
            "No JWT yet — user_id is trusted from the phone. Auth is a planned next step.",
            "Training pipeline runs offline; weights are pushed via the backend response. A scheduled retrain on Supabase data is the natural next milestone.",
        ],
    )

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
